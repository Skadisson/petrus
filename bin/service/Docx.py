from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from datetime import datetime
from bin.service import Environment
from bin.service import Cache
import matplotlib.pyplot as plt
import numpy
import math


class Docx:
    """DOCX generator"""

    def __init__(self):
        self.document = Document()
        self.environment = Environment.Environment()
        self.figure_number = 1
        self.cache = Cache.Cache()

    def place_headline(self):
        now = datetime.now()
        date = datetime.strftime(now, "%Y/%m/%d")
        self.document.add_heading("Support Report {}".format(date), 0)

    def place_stats(self, ticket_count, internal_count, external_count, hours_total, lifetime_per_ticket, hours_per_type, total_score, top_5_ticket_ranks, bottom_5_ticket_ranks, months):

        self.document.add_paragraph("").add_run("Durch Petrus automatisiert erstellt.").italic = True
        self.document.add_paragraph("").add_run("Petrus kennt und berechnet keine personen bezogenen Daten, alle angegebenen Rechnungen und Werte betreffen den gesamten Service Desk, das heißt reine CS-Tickets so wie auch auf Produkt-, QS-, DevOps- oder Projekt-Bords gesyncte Tickets.").italic = True

        lifetime_count = 0
        lifetime_total = 0.0
        lifetime_max = 0.0
        max_days_key = ''
        for lifetime in lifetime_per_ticket:
            if lifetime[1] > 0:
                lifetime_count += 1
                lifetime_total += lifetime[1]
                if lifetime[1] > lifetime_max:
                    lifetime_max = lifetime[1]
                    max_days_key = lifetime[0]
        if lifetime_count > 0:
            lifetime_average = (lifetime_total / lifetime_count)
        else:
            lifetime_average = 0.0

        if ticket_count > 0.0:
            average = hours_total/ticket_count
        else:
            average = 0.0
        hours_average = round(average, ndigits=2)

        lifetime_average_days = round(lifetime_average/24, ndigits=2)
        lifetime_max_days = round(lifetime_max/24, ndigits=2)

        support_relation, bugfix_relation = self.calculate_type_relation(hours_per_type)

        hours_total = round(hours_total)
        days = self.months_to_days(months)
        paragraph = self.document.add_paragraph("In den letzten ")
        paragraph.add_run("{} Tagen".format(days)).bold = True
        paragraph.add_run(" wurden ")
        paragraph.add_run("{} Tickets".format(ticket_count)).bold = True
        paragraph.add_run(" getrackt. Davon waren ")
        paragraph.add_run("{} Tickets von Mitarbeitern".format(internal_count)).bold = True
        paragraph.add_run(" und ")
        paragraph.add_run("{} Tickets von Kunden".format(external_count)).bold = True
        paragraph.add_run(". Insgesamter getrackter Aufwand war ")
        paragraph.add_run("{} Stunden".format(hours_total)).bold = True
        paragraph.add_run(". Fehler-Support-Verhältnis war ")
        paragraph.add_run("{}:{}".format(bugfix_relation, support_relation)).bold = True
        paragraph.add_run(". Durchschnittlicher Aufwand pro Ticket war damit ")
        paragraph.add_run("{} Stunden".format(hours_average)).bold = True
        paragraph.add_run(".")
        self.document.add_paragraph(f"Die durchschnittliche Wartezeit bis zur Lösung war {lifetime_average_days} Tage, die längste Wartezeit war {lifetime_max_days} Tage in Ticket {max_days_key}.")
        paragraph = self.document.add_paragraph(f"Die {ticket_count} Tickets haben eine Gesamt Straf-Punktezahl von ")
        paragraph.add_run(f"{total_score}").bold = True
        paragraph.add_run(" erhalten.")
        paragraph = self.document.add_paragraph("Davon sind die 5 Tickets, die am schlechtesten abgeschnitten haben:")
        for ticket_key in top_5_ticket_ranks:
            paragraph.add_run(f" {ticket_key}").bold = True
            paragraph.add_run(f" ({top_5_ticket_ranks[ticket_key]})")
        paragraph.add_run("; und die 5 Tickets, die am besten abgeschnitten haben:")
        for ticket_key in bottom_5_ticket_ranks:
            paragraph.add_run(f" {ticket_key}").bold = True
            paragraph.add_run(f" ({bottom_5_ticket_ranks[ticket_key]})")
        paragraph.add_run(".")

    @staticmethod
    def months_to_days(months):
        if months > 0:
            days = int(months * 30)
        else:
            days = 30

        return days

    def place_summary(self, _summary, _months):
        days = self.months_to_days(_months)
        paragraph = self.document.add_paragraph('')
        paragraph.add_run(f"Zusammengefasst wurden in den letzten {days} Tagen in Tickets folgende 'Problem', 'Lösung', 'Aufgabe' und 'Grund' Notizen hinterlegt:").bold = True
        self.document.add_paragraph(_summary)

    def place_projects(self, hours_per_project, project_ticket_count, hours_per_system, system_ticket_count, system_versions, months):
        days = self.months_to_days(months)
        self.document.add_heading('Projekte und Systeme', level=1)
        self.document.add_paragraph('Im folgenden getrackte Aufwände der letzten {} Tage pro Projekt.'.format(days))
        for project_hours in hours_per_project:
            paragraph = self.document.add_paragraph('')
            project_name = project_hours[0]
            hours = project_hours[1]
            paragraph.add_run("{}".format(project_name)).bold = True
            paragraph.add_run(" - {} Stunden auf {} Tickets".format(round(hours, ndigits=2), project_ticket_count[project_name]))
            if project_name in hours_per_system:
                for system_name in hours_per_system[project_name]:
                    paragraph = self.document.add_paragraph('')
                    system_hour_count = sum(hours_per_system[project_name][system_name])
                    paragraph.add_run("-> {}".format(system_name))
                    if project_name in system_versions and system_name in system_versions[project_name] and len(system_versions[project_name][system_name]) > 0:
                        paragraph.add_run(" - {} Stunden auf {} Tickets in {}".format(round(system_hour_count, ndigits=2),
                                                                                      system_ticket_count[project_name][system_name],
                                                                                      ", ".join(system_versions[project_name][
                                                                                                    system_name])))
                    else:
                        paragraph.add_run(" - {} Stunden auf {} Tickets".format(round(system_hour_count, ndigits=2),
                                                                                system_ticket_count[project_name][system_name]))

    def place_type_weight(self, hours_per_version, projects_per_version, months):
        bb_versions = self.environment.get_bb_versions()
        weights = {}
        projects = {}
        for bb_type in bb_versions:
            weights[bb_type] = 0.0
        days = self.months_to_days(months)
        self.document.add_heading('Gewichtung', level=1)
        self.document.add_paragraph('Folgende brandbox Typen haben in {} Tagen getrackte Aufwände erzeugt.'.format(days))
        for version_hours in hours_per_version:
            version = version_hours[0]
            hours = round(version_hours[1], ndigits=2)
            for bb_type in bb_versions:
                bb_type_versions = bb_versions[bb_type].split(" ")
                if version in bb_type_versions:
                    weights[bb_type] += hours
                    if version in projects_per_version:
                        if bb_type not in projects:
                            projects[bb_type] = []
                        for project in projects_per_version[version]:
                            if project not in projects[bb_type]:
                                projects[bb_type].append(project)
        for bb_type in weights:
            if bb_type in projects:
                label = bb_type
                if label == 'n/a':
                    label = 'nicht angegeben'
                bb_type_versions = bb_versions[bb_type].split(" ")
                paragraph = self.document.add_paragraph('')
                if bb_type_versions[0] != bb_type_versions[-1]:
                    paragraph.add_run("{} ({})".format(label, bb_type_versions[0] + ' - ' + bb_type_versions[-1])).bold = True
                else:
                    paragraph.add_run("{}".format(label)).bold = True
                paragraph.add_run(" - {} Stunden auf {} Projekte".format(weights[bb_type], len(projects[bb_type])))
        total = sum(list(weights.values()))
        labels = []
        for label in list(weights.keys()):
            if total > 0:
                sub_total = (weights[label] / total) * 100
            else:
                sub_total = 0
            labels.append(f"{label} [{str(round(sub_total))}%]")
        self.place_pie_chart(list(weights.values()), labels)
        for version_hours in hours_per_version:
            label = version_hours[0]
            hours = version_hours[1]
            if label == 'n/a':
                label = 'nicht angegeben'
            paragraph = self.document.add_paragraph('')
            paragraph.add_run("{}".format(label)).bold = True
            paragraph.add_run(" - {} Stunden".format(round(hours, ndigits=2)))

    def place_tickets(self, hours_per_ticket, months):
        days = self.months_to_days(months)
        self.document.add_heading('Bearbeitete SERVICE Tickets der letzten {} Tage mit Aufwänden'.format(days), level=1)
        for ticket_hours in hours_per_ticket:
            if ticket_hours[1] > 0.0:
                title = self.cache.get_ticket_title_by_key(ticket_hours[0])
                summary = self.cache.get_ticket_summary_by_key(ticket_hours[0])
                get_project_name_by_key = self.cache.get_project_name_by_key(ticket_hours[0])
                paragraph = self.document.add_paragraph('')
                paragraph.add_run("{}".format(ticket_hours[0])).bold = True
                if get_project_name_by_key != '':
                    paragraph.add_run(" - {} ".format(get_project_name_by_key))
                if title != '':
                    paragraph.add_run(" - {} ".format(title))
                paragraph.add_run(" - ")
                paragraph.add_run("{} Stunden".format(round(ticket_hours[1], ndigits=2))).bold = True
                if summary != '':
                    self.document.add_paragraph(summary)

    @staticmethod
    def calculate_type_relation(hours_per_type):

        support_hours = 0.0
        bugfix_hours = 0.0
        for type_hours in hours_per_type:
            type = type_hours[0]
            hours = type_hours[1]
            if type in ['Fehler', 'Bug', 'Maintenance']:
                bugfix_hours += hours
            else:
                support_hours += hours
        hours_sum = support_hours + bugfix_hours
        if hours_sum > 0.0:
            support_relation = round((support_hours / hours_sum) * 100)
            bugfix_relation = round((bugfix_hours / hours_sum) * 100)
        else:
            support_relation = 50
            bugfix_relation = 50

        return support_relation, bugfix_relation

    @staticmethod
    def calculate_hours_per_keyword_relation(hours_per_keyword):
        values = []
        labels = []
        i = 5
        for keyword_set in hours_per_keyword:
            values.append(keyword_set['hours'])
            labels.append(f"{keyword_set['keyword']} ({round(keyword_set['hours'])}h)")
            i -= 1
            if i == 0:
                break
        return values, labels

    @staticmethod
    def calculate_payed_unpayed_relation(payed_unpayed):
        values = []
        labels = []
        for label in payed_unpayed:
            values.append(payed_unpayed[label]['value'])
            labels.append(payed_unpayed[label]['label'])
        return values, labels

    def place_type_pie_chart(self, hours_per_type):

        support_relation, bugfix_relation = self.calculate_type_relation(hours_per_type)

        values = [support_relation, bugfix_relation]
        labels = [f"Support [{support_relation}%]", f"Fehler [{bugfix_relation}%]"]
        self.place_pie_chart(values, labels)

    def place_payed_unpayed_pie_chart(self, payed_unpayed, months):
        days = self.months_to_days(months)
        self.document.add_heading('Abrechnungspotenzial', level=1)
        total = math.ceil(payed_unpayed['payed']['value'] + payed_unpayed['unpayed']['value'])
        self.document.add_paragraph('Abrechenbare (Kostenübernahme = Kunde, Projekt, Abgerechnet) und nicht abrechenbare (Kostenübernahme = Konmedia, Leer) Aufwände der letzten {} Tage auf {} Stunden.'.format(days, total))
        values, labels = self.calculate_payed_unpayed_relation(payed_unpayed)
        self.place_pie_chart(values, labels)

    def place_pie_chart(self, values, labels):

        plot_path = self.environment.get_path_plot()
        sub_plot_path = str(plot_path).replace('plot', f"pie_plot_{self.figure_number}")

        plt.figure(self.figure_number)
        self.figure_number += 1
        colors = ['#00FFAE', '#FD5A2F', '#16BAE7']
        if len(values) > 3:
            colors = ['#16BAE7', '#4DCEF1', '#93E2F7', '#BFEEFB', '#E9FAFF']
        plt.pie(values, labels=labels, colors=colors)
        plt.savefig(sub_plot_path)

        self.document.add_picture(sub_plot_path, width=Inches(5))

    def place_page_break(self):
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)

    def place_plot(self, plot_data):
        plot_path = self.environment.get_path_plot()
        for axis in plot_data:
            sub_plot_path = str(plot_path).replace('plot', f"axis_plot_{self.figure_number}")

            axis_data = plot_data[axis]
            two_axis = axis.split('/')
            x_axis = two_axis[1]
            y_axis = two_axis[0]

            x_values = []
            y_values = []

            for x_value in axis_data:
                y_values.append(axis_data[x_value])
                if x_axis == 'day':
                    x_value = str(x_value)
                    x_value = str(f"{x_value[6]}{x_value[7]}.")
                elif x_axis == 'priority':
                    priorities = ["Unwesentlich", "Blocker", "Lowest", "Low", "Medium", "High", "Highest"]
                    x_value = priorities[(x_value+1)]
                x_values.append(x_value)

            y_avg = []
            if len(y_values) > 0:
                average_y = numpy.average(y_values)
                y_avg = [average_y]*len(x_values)

            plt.figure(self.figure_number)
            self.figure_number += 1
                
            plt.bar(x_values, y_values, color=(['#16BAE7']*len(y_values)))
            if len(y_avg) > 0:
                plt.plot(x_values, y_avg, color='#FD5A2F')

            if axis == "average days/priority":
                y_bonus = []
                for priority in x_values:
                    target_days = 7
                    if priority in ["Unwesentlich", "Blocker", "Lowest"]:
                        target_days = 30
                    elif priority in ["Low"]:
                        target_days = 21
                    elif priority in ["Medium"]:
                        target_days = 7
                    elif priority in ["High"]:
                        target_days = 2
                    elif priority in ["Highest"]:
                        target_days = 1
                    y_bonus.append(target_days)
                plt.ylim(0, 14)
                plt.bar(x_values, y_bonus, color='#00FFAE', alpha=.5)

            plt.xlabel(x_axis)
            plt.ylabel(y_axis)
            plt.grid(True)
            plt.savefig(sub_plot_path)

            headline = axis
            if axis == "new tickets/day":
                headline = "Neue minus geschlossene Tickets am Tag"
            elif axis == "average days/priority":
                headline = "Durchschnittliche Lebensdauer pro Priorität (ohne IS/PE Version)"
            self.document.add_heading(headline, level=1)
            self.document.add_picture(sub_plot_path, width=Inches(5))

    def save(self):
        path = 'temp/trend.docx'
        self.document.save(path)

        return path
